"""DOCX extractor using python-docx for text extraction."""

import logging
from pathlib import Path

from docx import Document

from .base import BaseExtractor, ExtractionResult

logger = logging.getLogger(__name__)


class DOCXExtractor(BaseExtractor):
    """Extractor for DOCX files using python-docx."""

    SUPPORTED_MIME_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    def extract(self, file_path: str) -> ExtractionResult:
        """
        Extract text and metadata from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ExtractionResult containing extracted text and metadata
        """
        try:
            path = Path(file_path)

            if not path.exists():
                return ExtractionResult(
                    text="",
                    metadata={},
                    success=False,
                    error=f"File not found: {file_path}",
                )

            # Open DOCX document
            doc = Document(file_path)

            # Extract text from all paragraphs
            text_parts = []
            paragraph_count = 0

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Only include non-empty paragraphs
                    text_parts.append(text)
                    paragraph_count += 1

            # Extract text from tables
            table_count = len(doc.tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            # Combine all text
            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                return ExtractionResult(
                    text="",
                    metadata=self._create_metadata(
                        word_count=0,
                        paragraph_count=0,
                    ),
                    success=False,
                    error="No text content found in DOCX",
                )

            word_count = self._count_words(full_text)

            metadata = self._create_metadata(
                word_count=word_count,
                paragraph_count=paragraph_count,
                table_count=table_count,
                file_size_bytes=path.stat().st_size,
            )

            logger.info(
                f"Successfully extracted text from DOCX {file_path}: "
                f"{paragraph_count} paragraphs, {table_count} tables, {word_count} words"
            )

            return ExtractionResult(
                text=full_text,
                metadata=metadata,
                success=True,
            )

        except Exception as e:
            logger.exception(f"Error extracting text from DOCX {file_path}: {e}")
            return ExtractionResult(
                text="",
                metadata={},
                success=False,
                error=str(e),
            )

    def supports_mime_type(self, mime_type: str) -> bool:
        """
        Check if this extractor supports the given MIME type.

        Args:
            mime_type: MIME type to check

        Returns:
            True if this extractor supports DOCX files
        """
        return mime_type in self.SUPPORTED_MIME_TYPES
